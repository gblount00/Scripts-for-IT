import tkinter as tk
from tkinter import ttk
import os
import subprocess
import webbrowser
from datetime import datetime
from docx import Document
ID_FILE = "current_id.txt"

# Function to create check-out tab content
def create_checkout_tab(notebook):
    checkout_frame = tk.Frame(notebook)
    
    # Example label and button for the check-out tab
    checkout_label = tk.Label(checkout_frame, text="Active Guest List")
    checkout_label.pack(pady=20)
    
    checkout_button = tk.Button(checkout_frame, text="Checkout Guest", font=("Arial", 14))
    checkout_button.pack(pady=10)
    
    # Add the frame to the notebook
    notebook.add(checkout_frame, text="Guest Check-out")

def Current_Id_Count():
    000000 

def update_time():
    current_datetime = datetime.now().strftime("%m/%d/%Y %H:%M:%S")
    Date_Time_Label.config(text=current_datetime)
    # Call this function again after 1000 ms (1 second)
    Date_Time_Label.after(1000, update_time)


# Function to increment the ID number
def increment_id():
    global current_id
    ID_Number.config(text=str(current_id))
    current_id += 1

# Function to save the current ID to a file
def save_current_id():
    with open(ID_FILE, 'w') as f:
        f.write(str(current_id))

# Function to load the current ID from a file
def load_current_id():
    global current_id
    if os.path.exists(ID_FILE):
        with open(ID_FILE, 'r') as f:
            current_id = int(f.read().strip())
    else:
        current_id = 555555  # Default starting ID if file doesn't exist


def print_info():
    # Gather data from entry fields
    guest_name = Guest_Name_Entry.get()
    id_number = ID_Number_Entry.get()
    date_time = Date_Time_Entry.get()
    
    # Create a Word document
    doc = Document()
    
    # Add content to the document
    doc.add_heading('Guest Check-In Information', level=1)
    doc.add_paragraph(f'Guest Name: {guest_name}')
    doc.add_paragraph(f'ID Number: {id_number}')
    doc.add_paragraph(f'Date & Time: {date_time}')
    
    # Save the document
    doc_file_path = 'guest_checkin_info.docx'
    doc.save(doc_file_path)
    
    # Print the document (this will use the default printer)
    os.startfile(doc_file_path, "print")

window = tk.Tk()
window.title("Guest Badge Creator")
window.geometry("500x600")

#the entire window
frame = tk.Frame(window)
frame.pack()


# Create a Notebook widget for tabs
notebook = ttk.Notebook(window)
notebook.pack(expand=True, fill='both')

# Create the Guest Check-in tab
checkin_frame = tk.Frame(notebook)
notebook.add(checkin_frame, text="Guest Check-in")



# Define consistent margin values
margin_x = 20  # Horizontal margin
margin_y = 5  # Vertical margin
widget_width = 40  # Width for Entry and ComboBox


#---------------GUEST INFO FRAME---------------
Guest_Check_in_Frame= tk.LabelFrame(checkin_frame, text="Guest Check-in")
Guest_Check_in_Frame.grid(row = 0, column = 0, padx=margin_x, pady=margin_y)

#guest name field
Guest_Name = tk.Label(Guest_Check_in_Frame, text="Guest Name: ")
Guest_Name.grid(row = 0, column = 0, sticky="w", padx=margin_x)
Guest_Name_Entry = tk.Entry(Guest_Check_in_Frame, width=widget_width)
Guest_Name_Entry.grid(row = 0, column = 1, sticky="we", padx=margin_x)

#category name field (visitor, contractor, etc)
Category_Name = tk.Label(Guest_Check_in_Frame, text="Category: ")
Category_Name.grid(row = 1, column = 0, sticky="w", padx=margin_x)
Category_Combobox = ttk.Combobox(Guest_Check_in_Frame, width=widget_width, value=["Visitor", "Visitor - Child", "Contractor","Vendor"])
Category_Combobox.grid(row = 1, column = 1, padx=margin_x, sticky="we")


#comapny name field
Company_Name =tk.Label(Guest_Check_in_Frame, text="Company Name: ")
Company_Name.grid(row = 2, column = 0, sticky="w", padx=margin_x)
Company_Name_Entry = tk.Entry(Guest_Check_in_Frame, width=widget_width)
Company_Name_Entry.grid(row = 2, column = 1, padx=margin_x, sticky="we")





# ---------LOCATION FRAME--------------
Location_Frame = tk.LabelFrame(checkin_frame, text="Visit Location: ")
Location_Frame.grid(row=3, column=0, padx=margin_x, pady=margin_y)

#building number combobox (selector)
Building_Number = tk.Label(Location_Frame, text = "Building Number: ")
Building_Number.grid(row = 3, column = 0, sticky="w", padx=margin_x)
Building_Combobox = ttk.Combobox(Location_Frame, width=widget_width, value=["4050 CCOC", "4030 CCOC"])
Building_Combobox.grid(row = 3, column = 1, sticky="we", padx=margin_x)

Agency_Visiting = tk.Label(Location_Frame, text="Agency Visiting: ")
Agency_Visiting.grid(row=4, column=0, padx=margin_x, sticky="w")
Agency_Visiting_Combobox = ttk.Combobox(Location_Frame, width=widget_width, value=["DMS Visitor", "DACS Visitor"])
Agency_Visiting_Combobox.grid(row=4, column=1, padx=margin_x)

Person_Checking_in = tk.Label(Location_Frame, text="Checked in By: ")
Person_Checking_in.grid(row=5, column=0, padx=margin_x, sticky="w")
Person_Checking_in_Combobox = ttk.Combobox(Location_Frame, width=widget_width, value=["Brenda Eatman", "John Case", "Kourtney Fields", "Latoria McKelvey", "Matthew Seeger", "Trinica Kirkland", "Other"])
Person_Checking_in_Combobox.grid(row=5, column=1, padx=margin_x, sticky="we")

Other_Person_Checking_in = tk.Label(Location_Frame, text="Other: ")
Person_Checking_in_Entry =tk.Entry(Location_Frame, width=widget_width)
Other_Person_Checking_in.grid(row=6, column=0, sticky="w", padx=margin_x)
Person_Checking_in_Entry.grid(row=6, column=1, sticky="we", padx=margin_x)


Other_Person_Checking_in.grid_remove()
Person_Checking_in_Entry.grid_remove()


def Person_Check_In_Change(event):
    if Person_Checking_in_Combobox.get() == "Other":
        Other_Person_Checking_in.grid()
        Person_Checking_in_Entry.grid()

    else:
        Other_Person_Checking_in.grid_remove()
        Person_Checking_in_Entry.grid_remove()

Person_Checking_in_Combobox.bind("<<ComboboxSelected>>", Person_Check_In_Change)


#------------------AUTO GENERATED FRAME--------------
Auto_Frame = tk.LabelFrame(checkin_frame, text="auto generated")
Auto_Frame.grid(row=8, column=0, padx=margin_x, pady=margin_y, sticky="we")

current_id = 555555

#id number, auto gentetated
ID_Label = tk.Label(Auto_Frame, text="ID: ")
ID_Label.grid(row = 8, column = 0, sticky="we", padx=margin_x)
ID_Number = tk.Label(Auto_Frame, text=str(current_id), width=20)
ID_Number.grid(row=8, column=1, sticky="we", padx=10)

#date&time auto genereated
Date_Time = tk.Label(Auto_Frame, text="Date & Time: ")
Date_Time.grid(row = 9, column = 0, sticky="we", padx=margin_x)
Date_Time_Label = tk.Label(Auto_Frame, text="", width=widget_width)
Date_Time_Label.grid(row=9, column=1, sticky="we", padx=margin_x)

update_time()

#--------print button frame----------------
Button_Frame = tk.Frame(checkin_frame, padx=margin_x, pady=margin_y)
Button_Frame.grid()
Print_Button = tk.Button(Button_Frame, text="Print",font=("Arial", 14), width=25, height =1, command=increment_id)
Print_Button.grid()



#customizing the entire grid
for widget in Guest_Check_in_Frame.winfo_children():
    widget.grid_configure(pady=margin_y)

for widget in Location_Frame.winfo_children():
    widget.grid_configure(pady=margin_y)

for widget in Auto_Frame.winfo_children():
    widget.grid_configure(pady=margin_y)


Other_Person_Checking_in.grid_remove()
Person_Checking_in_Entry.grid_remove()

create_checkout_tab(notebook)

window.mainloop()